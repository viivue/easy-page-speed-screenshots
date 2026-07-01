import os
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _docx_available = True
except ImportError:
    _docx_available = False

HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


def _epss_add_hyperlink(paragraph, url, text):
    r_id = paragraph.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1155CC"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _epss_result_cell(cell, desktop_link, mobile_link):
    """Fill the Result cell with Desktop and Mobile lines, hyperlinked if available."""
    p_desktop = cell.paragraphs[0]
    if desktop_link:
        _epss_add_hyperlink(p_desktop, desktop_link, "Desktop")
    else:
        p_desktop.add_run("Desktop")

    p_mobile = cell.add_paragraph()
    if mobile_link:
        _epss_add_hyperlink(p_mobile, mobile_link, "Mobile")
    else:
        p_mobile.add_run("Mobile")


def epss_build_docx_report(rows, out_path, report_date=None, dropbox_configured=True):
    """
    rows: list of {"url", "title", "desktop_link", "mobile_link"} — one entry per URL.
    Links may be None. Returns None if python-docx is not installed.
    """
    if not _docx_available:
        return None
    report_date = report_date or datetime.now().strftime("%Y-%m-%d")
    doc = Document()
    doc.add_heading("PageSpeed Insights Report", level=1)
    doc.add_paragraph(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).runs[0].italic = True
    if not dropbox_configured:
        doc.add_paragraph(
            "Dropbox was not configured for this run, so the result links are empty. "
            "Set the Dropbox credentials and re-run to populate links."
        )
    doc.add_paragraph(
        "Each row links to dated screenshots of the live result stored in Dropbox. "
        "The visible Google branding in the image is the tamper-evident record."
    )

    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"

    for i, h in enumerate(["URL", "Result"]):
        table.rows[0].cells[i].text = h

    for r in rows:
        c = table.add_row().cells
        url = r.get("url", "")
        title = r.get("title") or url
        if url:
            _epss_add_hyperlink(c[0].paragraphs[0], url, title)
        else:
            c[0].text = title
        _epss_result_cell(c[1], r.get("desktop_link"), r.get("mobile_link"))

    doc.save(out_path)
    return out_path
